from fastapi import APIRouter, Depends, HTTPException, Response
from fastapi.responses import StreamingResponse
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy import select
import io
from datetime import datetime

# PDF Generation
from reportlab.lib.pagesizes import LETTER
from reportlab.pdfgen import canvas
from reportlab.lib import colors
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle

# Word Generation
from docx import Document
from docx.shared import Inches, Pt

from db.database import get_db
from db.models import User, LearningSession, LearningStep
from auth.dependencies import get_current_user

router = APIRouter(prefix="/export", tags=["Export"])

@router.get("/session/{session_id}/pdf")
async def export_roadmap_pdf(
    session_id: int,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    # 1. Fetch Session and Steps
    result = await db.execute(
        select(LearningSession).where(LearningSession.id == session_id, LearningSession.user_id == current_user.id)
    )
    session = result.scalar_one_or_none()
    if not session:
        raise HTTPException(status_code=404, detail="Roadmap session not found")

    result = await db.execute(
        select(LearningStep).where(LearningStep.session_id == session_id).order_by(LearningStep.step_number.asc())
    )
    steps = result.scalars().all()

    # 2. Generate PDF using ReportLab
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=LETTER)
    styles = getSampleStyleSheet()
    
    # Custom Styles
    title_style = ParagraphStyle(
        'TitleStyle',
        parent=styles['Heading1'],
        fontSize=24,
        textColor=colors.HexColor("#2563eb"),
        spaceAfter=20
    )
    
    day_style = ParagraphStyle(
        'DayStyle',
        parent=styles['Heading2'],
        fontSize=16,
        textColor=colors.HexColor("#1e293b"),
        spaceBefore=12,
        spaceAfter=6
    )

    elements = []
    
    # Header
    elements.append(Paragraph(f"Learning Roadmap: {session.topic}", title_style))
    elements.append(Paragraph(f"Created: {datetime.utcnow().strftime('%Y-%m-%d')}", styles['Normal']))
    elements.append(Paragraph(f"Duration: {session.total_steps} day(s)", styles['Normal']))
    elements.append(Spacer(1, 20))

    # Steps
    for step in steps:
        elements.append(Paragraph(f"Day {step.step_number}: {step.title}", day_style))
        # Handle markdown-like content (simplified for PDF)
        content_lines = step.content.split('\n')
        for line in content_lines:
            if line.strip():
                elements.append(Paragraph(line.strip(), styles['Normal']))
        elements.append(Spacer(1, 10))

    doc.build(elements)
    buffer.seek(0)
    
    filename = f"Roadmap_{session.topic.replace(' ', '_')}.pdf"
    return StreamingResponse(
        buffer,
        media_type="application/pdf",
        headers={"Content-Disposition": f"attachment; filename={filename}"}
    )

@router.get("/session/{session_id}/docx")
async def export_roadmap_docx(
    session_id: int,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    # 1. Fetch Session and Steps
    result = await db.execute(
        select(LearningSession).where(LearningSession.id == session_id, LearningSession.user_id == current_user.id)
    )
    session = result.scalar_one_or_none()
    if not session:
        raise HTTPException(status_code=404, detail="Roadmap session not found")

    result = await db.execute(
        select(LearningStep).where(LearningStep.session_id == session_id).order_by(LearningStep.step_number.asc())
    )
    steps = result.scalars().all()

    # 2. Generate Word using python-docx
    doc = Document()
    
    # Title
    title = doc.add_heading(f"Learning Roadmap: {session.topic}", 0)
    
    # Meta Info
    doc.add_paragraph(f"Created: {datetime.utcnow().strftime('%Y-%m-%d')}")
    doc.add_paragraph(f"Duration: {session.total_steps} day(s)")
    doc.add_paragraph("-" * 40)

    # Steps
    for step in steps:
        h2 = doc.add_heading(f"Day {step.step_number}: {step.title}", level=1)
        doc.add_paragraph(step.content)
        doc.add_paragraph("") # Space between days

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    filename = f"Roadmap_{session.topic.replace(' ', '_')}.docx"
    return StreamingResponse(
        buffer,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename={filename}"}
    )
